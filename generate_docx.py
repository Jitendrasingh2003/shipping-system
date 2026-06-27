import os
import re
import sys

def install_and_import(package, import_name):
    import importlib
    try:
        importlib.import_module(import_name)
    except ImportError:
        import subprocess
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Ensure python-docx is installed
install_and_import('python-docx', 'docx')

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for table cells in dxas (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_type, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_type}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_inline_formatting(paragraph, text, font_name="Times New Roman", font_size=12, italic_default=False):
    """Parse basic markdown inline formatting like **bold** and *italic* and add runs."""
    # Split text into parts by bold markdown **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        is_bold = False
        is_italic = italic_default
        clean_part = part
        
        if part.startswith('**') and part.endswith('**'):
            is_bold = True
            clean_part = part[2:-2]
            
        # Parse italic in the clean part
        subparts = re.split(r'(\*.*?\*|_.*?_)', clean_part)
        for subpart in subparts:
            sub_bold = is_bold
            sub_italic = is_italic
            clean_subpart = subpart
            
            if (subpart.startswith('*') and subpart.endswith('*')) or (subpart.startswith('_') and subpart.endswith('_')):
                sub_italic = True
                clean_subpart = subpart[1:-1]
                
            if clean_subpart:
                run = paragraph.add_run(clean_subpart)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = sub_bold
                run.italic = sub_italic
                # Keep font name working for non-ascii
                rPr = run._r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rPr.append(rFonts)

def main():
    md_file = "SmartShip_Project_Report.md"
    docx_file = "SmartShip_Project_Report.docx"
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found in current directory.")
        return

    print("Reading Markdown file...")
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Configure 1-inch margins (standard JECRC report margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    print("Compiling Word Document styles...")
    
    # Track states
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 1. Code Block Handler
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                # Create a single-cell table for shaded box
                code_table = doc.add_table(rows=1, cols=1)
                code_table.autofit = False
                code_table.columns[0].width = Inches(6.5)
                cell = code_table.cell(0, 0)
                set_cell_background(cell, "F1F5F9")  # light gray
                set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
                
                # Add code text
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.05
                p.paragraph_format.space_after = Pt(0)
                
                code_text = "\n".join(code_content)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(15, 23, 42)
                
                # Add spacing after code block
                doc.add_paragraph().paragraph_format.space_before = Pt(6)
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\n'))
            i += 1
            continue

        # 2. Page Break Handler
        if stripped == '\\pagebreak':
            doc.add_page_break()
            i += 1
            continue

        # 3. Horizontal Rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            # Add line
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'cbd5e1')
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # 4. Table Handler
        if stripped.startswith('|'):
            in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            # We reached a line that is not a table row
            # Render the accumulated table
            if table_rows:
                # Clean up delimiter rows like | :--- | :--- |
                cleaned_rows = []
                for row in table_rows:
                    if re.match(r'^\|\s*[:\-|\s]+\s*\|$', row):
                        continue
                    cleaned_rows.append(row)
                
                if cleaned_rows:
                    # Parse rows and cell values
                    parsed_grid = []
                    for row in cleaned_rows:
                        # Split by '|' and strip spaces, ignore leading/trailing empty elements
                        cells = [c.strip() for c in row.split('|')]
                        if len(cells) > 1:
                            # If row starts and ends with |, cells[0] and cells[-1] are empty
                            if cells[0] == '':
                                cells = cells[1:]
                            if cells[-1] == '':
                                cells = cells[:-1]
                            parsed_grid.append(cells)
                    
                    if parsed_grid:
                        num_cols = max(len(row) for row in parsed_grid)
                        word_table = doc.add_table(rows=len(parsed_grid), cols=num_cols)
                        word_table.style = 'Table Grid'
                        
                        for row_idx, row_data in enumerate(parsed_grid):
                            w_row = word_table.rows[row_idx]
                            is_header = (row_idx == 0)
                            
                            for col_idx in range(num_cols):
                                if col_idx < len(row_data):
                                    cell_val = row_data[col_idx]
                                else:
                                    cell_val = ""
                                    
                                cell = w_row.cells[col_idx]
                                p = cell.paragraphs[0]
                                p.paragraph_format.space_before = Pt(4)
                                p.paragraph_format.space_after = Pt(4)
                                
                                # Formatting headers
                                if is_header:
                                    set_cell_background(cell, "1E293B")  # Navy dark header
                                    set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
                                    run = p.add_run(cell_val)
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(10.5)
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                else:
                                    # Alternating row background shading
                                    if row_idx % 2 == 1:
                                        set_cell_background(cell, "F8FAFC") # soft shade
                                    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                                    add_inline_formatting(p, cell_val, font_size=10)
                                    
            in_table = False
            table_rows = []
            # Do not increment i, let the current line be processed in next iteration

        # 5. Heading 1
        if stripped.startswith('# '):
            heading_text = stripped[2:]
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            
            run = h.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59) # Slate Navy
            
            # Apply JECRC Center Sizing for cover layout
            if "SMART SHIP" in heading_text or "Industrial Project" in heading_text or "JECRC UNIVERSITY" in heading_text:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(20)
                
            i += 1
            continue

        # 6. Heading 2
        if stripped.startswith('## '):
            heading_text = stripped[3:]
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            
            run = h.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(15, 118, 110) # Teal
            i += 1
            continue

        # 7. Heading 3
        if stripped.startswith('### '):
            heading_text = stripped[4:]
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            
            run = h.add_run(heading_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85) # Medium Slate
            i += 1
            continue

        # 8. Bullets
        if stripped.startswith('* ') or stripped.startswith('- '):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            add_inline_formatting(p, bullet_text)
            i += 1
            continue

        # 9. Blockquotes
        if stripped.startswith('> '):
            quote_text = stripped[2:]
            quote_table = doc.add_table(rows=1, cols=1)
            quote_table.autofit = False
            quote_table.columns[0].width = Inches(6.5)
            cell = quote_table.cell(0, 0)
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # Left border thick gray
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '24')  # 3pt thickness
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), '0F766E') # Teal border
            tcBorders.append(left_border)
            
            # Clear other borders
            for border_name in ['top', 'bottom', 'right']:
                b = OxmlElement(f'w:{border_name}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)
            
            p = cell.paragraphs[0]
            add_inline_formatting(p, quote_text, italic_default=True)
            doc.add_paragraph().paragraph_format.space_before = Pt(4)
            i += 1
            continue

        # 10. Default Paragraphs
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            
            # Center alignment check for cover page tags
            if "[Logo" in stripped or "[Student Name]" in stripped or "[Faculty Guide" in stripped or "[Industry Guide" in stripped or "June 2024" in stripped or "Department of Computer" in stripped or "JECRC UNIVERSITY" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_formatting(p, stripped, font_size=13)
            elif "Signature of student" in stripped or "___________________________" in stripped or "Faculty Internship Guide" in stripped or "Industry Guide" in stripped or "Signature" in stripped:
                # Layout helpers for signature blocks (often left/right or right aligned)
                p.paragraph_format.space_before = Pt(20)
                add_inline_formatting(p, stripped, font_size=12)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_inline_formatting(p, stripped)
        else:
            # Empty line, add a minor spacing if needed
            pass
            
        i += 1

    print(f"Saving Document to {docx_file}...")
    doc.save(docx_file)
    print("Done! Word Document generated successfully.")

if __name__ == "__main__":
    main()
